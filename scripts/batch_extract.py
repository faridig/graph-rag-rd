#!/usr/bin/env python3
"""Batch experiment extractor — xlsx, docx, csv → 4 KG artifacts via Anthropic Sonnet 4.6.

Usage:
    python scripts/batch_extract.py                   # process all remaining files
    python scripts/batch_extract.py --dry-run         # show what would be processed
    python scripts/batch_extract.py --file NAME.xlsx  # process one specific file
    python scripts/batch_extract.py --force           # reprocess even if output exists

Output per file: lien_essai/{stem}/{id}_knowledge.json + _triples.csv + _documentation.md + _validation.md
"""
from __future__ import annotations

import argparse
import json
import logging
import os
import re
import subprocess
import sys
import time
from dataclasses import dataclass, field
from datetime import date
from pathlib import Path

# Ensure project root is on sys.path so `src.*` imports work
# when the script is invoked as `python scripts/batch_extract.py`
_PROJECT_ROOT = Path(__file__).resolve().parent.parent
if str(_PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(_PROJECT_ROOT))

import anthropic
import httpx
from dotenv import load_dotenv

load_dotenv()

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s %(levelname)s %(message)s",
    datefmt="%H:%M:%S",
)
log = logging.getLogger(__name__)
TODAY = date.today().isoformat()  # injected into provenance after LLM call

# ---------------------------------------------------------------------------
BRUT_DIR = Path("data/repertoire_rd_2025-2026/lien_essai/lien_essai_brut")
OUT_DIR = Path("data/repertoire_rd_2025-2026/lien_essai")
SCRIPTS_DIR = Path("scripts")
MODEL = "claude-sonnet-4-6"
MAX_TOKENS = 128_000         # output cap — Sonnet 4.6 max (128K)
THINKING_BUDGET = 10_000     # extended thinking budget (included in MAX_TOKENS)
MAX_INVENTORY_CHARS = 800_000
MAX_RETRIES = 2
SLEEP_BETWEEN = 2.0          # Sonnet 4.6 has higher rate limits than Opus 4.8

# ---------------------------------------------------------------------------
# Cost tracking — Anthropic pricing (USD per million tokens)
# Source: https://platform.claude.com/docs/en/about-claude/pricing (verified 2026-06-03)
# claude-sonnet-4-6: $3 input / $15 output / $3.75 cache_write_5m / $0.30 cache_read
# Extended thinking tokens billed at output rate (included in output_tokens usage field).
# Update if pricing changes: https://www.anthropic.com/pricing

_PRICING: dict[str, float] = {
    "input":        3.00,   # base input tokens
    "output":      15.00,   # output tokens (incl. thinking tokens)
    "cache_write":  3.75,   # 5-min cache write (1.25× input, ttl="5m")
    "cache_read":   0.30,   # cache read/hit (0.10× input)
}

# Batch API alternative (50% discount, async, hours delay):
# input=$1.50  output=$7.50 — use --batch flag if cost matters more than speed
_PRICING_BATCH: dict[str, float] = {
    "input":  1.50,
    "output":  7.50,
    "cache_write": 1.875,
    "cache_read":  0.15,
}


@dataclass
class _Cost:
    input_tokens:       int = 0
    output_tokens:      int = 0
    cache_write_tokens: int = 0
    cache_read_tokens:  int = 0

    def add_usage(self, usage: object) -> None:
        self.input_tokens       += getattr(usage, "input_tokens", 0) or 0
        self.output_tokens      += getattr(usage, "output_tokens", 0) or 0
        self.cache_write_tokens += getattr(usage, "cache_creation_input_tokens", 0) or 0
        self.cache_read_tokens  += getattr(usage, "cache_read_input_tokens", 0) or 0

    @property
    def usd(self) -> float:
        M = 1_000_000
        return (
            self.input_tokens       / M * _PRICING["input"]
            + self.output_tokens    / M * _PRICING["output"]
            + self.cache_write_tokens / M * _PRICING["cache_write"]
            + self.cache_read_tokens  / M * _PRICING["cache_read"]
        )

    def __add__(self, other: "_Cost") -> "_Cost":
        return _Cost(
            self.input_tokens       + other.input_tokens,
            self.output_tokens      + other.output_tokens,
            self.cache_write_tokens + other.cache_write_tokens,
            self.cache_read_tokens  + other.cache_read_tokens,
        )

    def summary(self) -> str:
        return (
            f"in={self.input_tokens:,} out={self.output_tokens:,} "
            f"cache_write={self.cache_write_tokens:,} cache_read={self.cache_read_tokens:,} "
            f"→ ${self.usd:.4f}"
        )

# Filesystem-unsafe characters in addition to whitespace
_UNSAFE_CHARS = re.compile(r'[/\\:*?"<>|]')

# ---------------------------------------------------------------------------
# Complexity triage
#
# Fichiers Niveau 3 : mesures instrumentales, réplicats TPA, profils thermiques,
# strides non-évidents → précision critique, skill interactif obligatoire.
# Le batch refuse de les traiter par défaut (--force-complex pour passer outre).

_COMPLEX_PREFIXES = (
    "ACE-1", "ACE-2", "ACE-4", "ACE-6",
    "FIB-1", "FIB-2", "FIB-3", "FIB-4", "FIB-10", "FIB-12",
    "STRIP-", "FIPROVEX-", "PP-REC-",
)

# Fichiers Niveau 2 : batch autorisé, mais _validation.md doit être relu par un humain.
_REVIEW_REQUIRED_PREFIXES = (
    "DST-", "GLU-", "Essais MDD", "Essais nuggets",
    "Essais transferts",
)


def _complexity_level(filename: str) -> int:
    """Return 1 (simple), 2 (review needed), or 3 (skill required)."""
    name = filename.upper()
    for p in _COMPLEX_PREFIXES:
        if name.startswith(p.upper()):
            return 3
    for p in _REVIEW_REQUIRED_PREFIXES:
        if name.upper().startswith(p.upper()):
            return 2
    return 1


# ---------------------------------------------------------------------------
# Skip-set detection

@dataclass
class SkipSet:
    source_files: set[str] = field(default_factory=set)
    exp_ids: set[str] = field(default_factory=set)


def _build_skip_set() -> SkipSet:
    """Collect source_file and experiment.id from every existing knowledge JSON.

    A brut file is considered already done if:
    - Its exact name (or underscore-normalised form) matches a source_file field, OR
    - Its name starts with an existing experiment.id followed by a separator char.
    """
    result = SkipSet()
    for kj in OUT_DIR.glob("**/*_knowledge.json"):
        try:
            data = json.loads(kj.read_text(encoding="utf-8"))
            exp = data.get("experiment", {})
            sf = exp.get("source_file", "")
            eid = exp.get("id", "")
            if sf and sf != "?":
                result.source_files.add(sf)
                result.source_files.add(sf.replace("_", " "))
            if eid and eid != "?":
                result.exp_ids.add(eid.strip())
        except Exception as e:
            log.debug("Could not read skip info from %s: %s", kj, e)
    return result


def _is_already_done(stem: str, force: bool) -> bool:
    if force:
        return False
    target = OUT_DIR / stem
    return target.is_dir() and any(target.glob("*_knowledge.json"))


def _file_matches_exp_id(filename: str, exp_ids: set[str]) -> bool:
    """True if filename starts with a known experiment ID followed by a separator."""
    for eid in exp_ids:
        pattern = re.escape(eid) + r"[-_. ]"
        if re.match(pattern, filename) or filename == eid:
            return True
    return False


# ---------------------------------------------------------------------------
# Onglet hint extraction

_ONGLET_PATTERNS = [
    re.compile(r'onglet\s+"([^"]+)"', re.IGNORECASE),   # onglet "Name"
    re.compile(r'\(onglet\s+([^)]+)\)', re.IGNORECASE),  # (onglet Name)
    re.compile(r'onglet\s+([A-Za-zÀ-ÿ0-9 _\-]+?)(?:\.xlsx|\.xlsm|$)', re.IGNORECASE),
]


def _parse_target_sheet(filename: str) -> str | None:
    """Extract the target sheet name from a filename containing 'onglet ...'."""
    for pat in _ONGLET_PATTERNS:
        m = pat.search(filename)
        if m:
            return m.group(1).strip().strip('"')
    return None


def _resolve_sheet(sheet_names: list[str], hint: str) -> str | None:
    """Case-insensitive match of hint against available sheet names.

    Priority: exact match > shortest partial match (most specific first).
    """
    hint_lower = hint.lower()
    for name in sheet_names:
        if name.lower() == hint_lower:
            return name
    # Partial match: collect all matches, return the shortest (most specific)
    partials = [
        name for name in sheet_names
        if hint_lower in name.lower() or name.lower() in hint_lower
    ]
    return min(partials, key=len) if partials else None


# ---------------------------------------------------------------------------
# Inventory helpers

def _dump_sheet(ws: object, lines: list[str], limit: int, total: int) -> int:  # type: ignore[type-arg]
    """Stream rows of ws into lines via iter_rows(values_only=True).

    Uses reset_dimensions() first so files with missing/corrupt dimension headers
    (common with read_only mode) are handled correctly — the official openpyxl fix.
    Returns updated total char count.
    """
    if hasattr(ws, "reset_dimensions"):  # ReadOnlyWorksheet only
        ws.reset_dimensions()  # type: ignore[attr-defined]
    for row_num, row in enumerate(ws.iter_rows(values_only=True), 1):  # type: ignore[attr-defined]
        if not any(c is not None and str(c).strip() for c in row):
            continue
        row_str = f"{row_num} || " + " | ".join("" if c is None else str(c) for c in row)
        lines.append(row_str)
        total += len(row_str)
        if total > limit:
            lines.append(f"... [sheet {ws.title!r} truncated]")  # type: ignore[attr-defined]
            break
    return total


def _inventory_xlsx(path: Path, target_sheet: str | None = None, limit: int = MAX_INVENTORY_CHARS) -> str:
    import openpyxl  # noqa: PLC0415

    wb = openpyxl.load_workbook(str(path), data_only=True, read_only=True)
    try:
        all_names = wb.sheetnames
        lines: list[str] = [f"WORKBOOK: {path.name}", f"SHEETS: {all_names}"]
        total = sum(len(l) for l in lines)

        if target_sheet:
            resolved = _resolve_sheet(all_names, target_sheet)
            if resolved:
                ws = wb[resolved]
                lines.append(f"\n[Focusing on sheet '{resolved}' as indicated by filename]")
                lines.append(f"\n=== SHEET: {resolved} ===")
                total = sum(len(l) for l in lines)
                total = _dump_sheet(ws, lines, limit, total)
            else:
                lines.append(f"\n[WARNING: sheet '{target_sheet}' not found — dumping all sheets]")
                target_sheet = None  # fall through to full dump

        if not target_sheet:
            for ws in wb.worksheets:
                header = f"\n=== SHEET: {ws.title} ==="
                lines.append(header)
                total += len(header)
                if total > limit:
                    lines.append("... [inventory truncated — too large]")
                    break
                total = _dump_sheet(ws, lines, limit, total)
    finally:
        wb.close()

    return "\n".join(lines)


def _inventory_csv(path: Path) -> str:
    import csv  # noqa: PLC0415

    lines = [f"CSV: {path.name}"]
    with open(path, newline="", encoding="utf-8", errors="replace") as f:
        try:
            dialect = csv.Sniffer().sniff(f.read(4096))
            f.seek(0)
        except Exception:
            dialect = csv.excel
            f.seek(0)
        for i, row in enumerate(csv.reader(f, dialect), 1):
            if any(c.strip() for c in row):
                lines.append(f"{i} || " + " | ".join(row))
    return "\n".join(lines)


def _inventory_docx(path: Path, limit: int = MAX_INVENTORY_CHARS) -> str:
    from docx import Document  # noqa: PLC0415

    doc = Document(str(path))
    lines: list[str] = [f"WORD DOCUMENT: {path.name}", ""]

    for para in doc.paragraphs:
        txt = para.text.strip()
        if not txt:
            continue
        # Prefix style name so LLM can distinguish titles from body text
        style = para.style.name if para.style else "Normal"
        if any(s in style for s in ("Heading", "Titre", "Title", "Subtitle")):
            lines.append(f"[{style}] {txt}")
        else:
            lines.append(txt)

    lines.append("\n--- TABLES ---")
    for ti, tbl in enumerate(doc.tables, 1):
        lines.append(f"\nTable {ti}:")
        for row in tbl.rows:
            seen: list[str] = []
            for cell in row.cells:
                txt = cell.text.strip()
                if not seen or seen[-1] != txt:
                    seen.append(txt)
            if any(seen):
                lines.append(" | ".join(seen))

    full = "\n".join(lines)
    if len(full) > limit:
        full = full[:limit] + "\n... [truncated]"
    return full


def get_inventory(path: Path) -> tuple[str, str | None]:
    """Return (inventory_text, target_sheet_hint)."""
    ext = path.suffix.lower()
    if ext in (".xlsx", ".xlsm"):
        target_sheet = _parse_target_sheet(path.name)
        return _inventory_xlsx(path, target_sheet=target_sheet), target_sheet
    if ext == ".csv":
        return _inventory_csv(path), None
    if ext == ".docx":
        return _inventory_docx(path), None
    raise ValueError(f"Unsupported extension: {ext}")


# ---------------------------------------------------------------------------
# Prompt (stable part cached, variable part in user message)

_SYSTEM_PROMPT = """\
You are an expert R&D knowledge engineer at a French food-tech company specialising in \
plant-based meat analogues (analogues de viande végétaux). You extract experiment files \
into canonical JSON for a knowledge graph used daily by R&D engineers to find, compare, \
and reproduce experiments.

## Domain context

Products: P01=poulet allumettes HME, P02=poulet émincés HME, M03=milanaise HME, \
B08=boulette, P16=poulet dés, M02=milanaise burger. Codes appear everywhere — always \
expand them in the glossary.

Equipment: EV32=twin-screw extruder bi-vis (HME), RVA=Rapid Visco Analyser, \
TPA=Texture Profile Analysis, SME=Specific Mechanical Energy (Wh/kg).

Key measurements and their units:
- TPA: hardness (g), springiness (ratio), cohesiveness (ratio), chewiness (g)
- Cut_T/Cut_L: transverse/longitudinal cutting force (g); anisotropy = Cut_T/Cut_L (target > 1.2)
- Barrel zones Z1–Z8 (°C), die pressure (bar), screw speed (rpm), throughput (kg/h)
- Moisture (%), water activity (aw), pH, colour L*a*b*, protein content (%)

Chantiers: Extrusion, Applications, MDD, Quick, Kobé, STRIP, DST, FIB, ACE, GLU, PP, VEILLE

## Target JSON schema

```
{
  "experiment": {
    "id": "ACE-5",                        // short, filesystem-safe
    "title": "...", "type": "...",
    "objective": "...", "date": "YYYY-MM-DD",
    "operator": "...", "equipment": "...",
    "scale": "lab|pilot|industrial",      // NEW: labo/pilote/industriel
    "batch_size": {"value": 4, "unit": "kg"},  // NEW
    "domain": "plant-based meat analogue",
    "source_file": "original_filename.xlsx",
    "status": "preliminary|ongoing|complete"  // NEW: see rule 13
  },
  "targets": {                            // NEW: acceptance criteria / objectifs
    "metric_name": {"min": 1.2, "max": null, "unit": null, "note": "objectif anisotropie"}
  },
  "references": ["ACE-3", "ACE-4"],       // NEW: cross-experiment links found in the file
  "glossary": {"TERM": "definition"},
  "design": {
    "factors": [{"name": "...", "unit": "...", "levels": [...]}],
    "control": "run_id", "note": "..."
  },
  "runs": [...],                          // valid runs only — see below
  "failed_runs": [                        // NEW: runs that did not complete
    {"id": "3", "name": "...", "failure_reason": "bouchage filière",
     "conditions_at_failure": {...}, "notes": "..."}
  ],
  "derived": [                            // extended: any computed quantity
    {"run": "2", "label": "...",
     "vs_control_pct": {"hardness": -14.0},   // % change vs control
     "computed": {"sme_wh_kg": 185.3, "protein_dry_basis_pct": 72.1}}  // NEW
  ],
  "observations": {
    "process": "...", "sensory": "...",   // NEW: sensory/organoleptic separate from process
    "conclusion": "...", "next_step": "..."
  },
  "not_measured": [
    {"analysis": "colorimetry", "reason": "planned but not done — time constraint"}
  ],                                      // CHANGED: object with reason, not just string
  "unused_palette": ["ingredient in template but not used"],
  "provenance": {"extraction_method": "batch_extract.py LLM", "generated_on": "YYYY-MM-DD"}
}
```

Each valid run:
```
{
  "id": "1", "name": "P02 + 0.2% NaCl", "is_control": false,
  "factor_levels": {"salt_type": "NaCl", "salt_dose": 0.2},
  "inputs": {
    "formulation": [
      {"component": "Nutralys F85M",
       "supplier": "Roquette",           // NEW: if present in file
       "lot_number": "B-1234",           // NEW: if present
       "role": "pea protein isolate",
       "pct_matrix": {"value": 70, "unit": "%"},
       "kg": {"value": 7.0, "unit": "kg"}}
    ]
  },
  "conditions": {"screw_speed": {"value": 625, "unit": "rpm"}, ...},
  "responses": {"cut_T": {"mean": 19470, "sd": 1069, "unit": "g", "replicates": [...]}, ...},
  "notes": "..."
}
```

## Extraction rules — R&D precision

1. **Setpoint AND actual** for every process parameter — both, never just one.
2. **Replicates alongside aggregates** — keep raw measurements, not just mean/SD.
3. **Ingredient traceability** — capture supplier and lot number when present anywhere in the file \
(side notes, header block, MP table).
4. **Failed runs are data** — a run labelled "STOP", "bouchage", "non exploitable", "échec" goes \
into `failed_runs` with its `failure_reason` and the conditions at the time of failure. \
Never discard them — they define process boundaries.
5. **Cross-experiment references** — if the file mentions "voir ACE-4", "suite de STRIP-3", \
"confirme résultat de FIB-2", capture the referenced ID in `references`.
6. **Acceptance criteria** — if the file states a target (">1.2", "objectif 25-35 kN"), \
capture it in `targets`.
7. **Sensory observations separate** — put organoleptic notes (texture, couleur, goût, aspect) \
in `observations.sensory`, not mixed with process observations.
8. **Scale and batch size** — always capture labo/pilote/industriel and kg produced.
9. **Computed quantities** — SME, rendement, formulation sur base sèche, etc. go in \
`derived[].computed`, flagged as calculated values.
10. **not_measured as objects** — always include `reason` (planned-but-skipped vs unavailable).
11. Missing data → null + note. Anomalies → note field. Never invent, never smooth over.
12. Glossary entry for every abbreviation, product code, and supplier name found in the file.
13. **Experiment status** — always set `experiment.status`:
    - "complete" if all planned measurements are done and conclusions are drawn
    - "preliminary" if some measurements are missing but partial conclusions exist
    - "ongoing" if the study is explicitly still in progress (Phase 2 en cours, résultats attendus, suite prévue)
14. **`vs_control_pct` contains only numeric percentages** — keys must map to numbers or null. \
Never put strings (notes, explanations) inside `vs_control_pct`; put them in the `note` field \
at the root of the `derived` entry instead.
15. **Units on every numeric measurement dict** — always include the `unit` key, even when \
the unit is already encoded in the field name. Standard mappings to apply systematically: \
`screw_speed` → `"rpm"`, `*_pct` or `couple_pct` → `"%"`, `*_bar` → `"bar"`, \
`*_C` or `tmatiere_C` → `"°C"`, `throughput_kg_h` → `"kg/h"`, `sme_wh_kg` → `"Wh/kg"`, \
`cut_L` / `cut_T` / `hardness` → `"g"`, `moisture` → `"%"`. \
Use `null` for truly dimensionless ratios (anisotropy, pH). Never omit the `unit` key.

⚠ CRITICAL — SELF-CONTAINED RUNS:
Every run MUST be fully self-contained. NEVER write "idem ESSAI X", "même formule", \
"cf. run Y", or any reference to another run's data. \
If run 2 has the same formulation as run 1, copy every ingredient with its exact quantity. \
A run with null `pct_matrix.value` is INCOMPLETE and causes data loss. \
A colleague must be able to reproduce any single run from its JSON entry alone, \
without reading any other run.
"""


def _build_user_message(
    filename: str, file_type: str, inventory: str, target_sheet: str | None = None
) -> str:
    sheet_note = (
        f"\nTarget sheet (from filename): **{target_sheet}** — extract only this sheet's data."
        if target_sheet
        else ""
    )
    return f"""\
## File to extract

Filename: {filename}
File type: {file_type}{sheet_note}

## Content dump

```
{inventory}
```

## Task

Generate a complete, valid `_knowledge.json`.
- Infer experiment.id from the filename when not explicit \
(e.g. "STRIP-10 Essais de coupe.xlsx" → id = "STRIP-10").
- If a target sheet is specified, base the extraction on that sheet only; \
use the sheet name to infer the experiment id when the filename has no explicit id \
(e.g. "Essais MDD - onglet \\"Kefta\\"" → id = "ESSAIS-MDD-KEFTA").
- Set `experiment.source_file` = the original filename (with extension).
- For docx: one run per discrete trial/fabrication described.

Return ONLY the JSON, inside ```json ... ``` fences. No explanation before or after.
"""


# ---------------------------------------------------------------------------
# API call with prompt caching

def _call_llm(user_message: str, client: anthropic.Anthropic) -> tuple[str, _Cost]:
    """Returns (response_text, cost_for_this_call). Uses streaming (required for max_tokens>10m)."""
    call_cost = _Cost()
    for attempt in range(MAX_RETRIES + 1):
        try:
            with client.messages.stream(
                model=MODEL,
                max_tokens=MAX_TOKENS,
                thinking={"type": "enabled", "budget_tokens": THINKING_BUDGET},
                system=[
                    {
                        "type": "text",
                        "text": _SYSTEM_PROMPT,
                        "cache_control": {"type": "ephemeral", "ttl": "5m"},
                    }
                ],
                messages=[{"role": "user", "content": user_message}],
            ) as stream:
                text = stream.get_final_text()
                msg = stream.get_final_message()

            usage = msg.usage
            call_cost.add_usage(usage)
            cache_read = getattr(usage, "cache_read_input_tokens", 0) or 0
            thinking_tokens = sum(
                len(b.thinking) // 4
                for b in msg.content
                if getattr(b, "type", "") == "thinking"
            )
            log.info(
                "  tokens in=%d out=%d thinking≈%d cache=%s write=%d read=%d → $%.4f",
                usage.input_tokens,
                usage.output_tokens,
                thinking_tokens,
                "HIT" if cache_read > 0 else "MISS",
                getattr(usage, "cache_creation_input_tokens", 0) or 0,
                cache_read,
                call_cost.usd,
            )
            return text, call_cost

        except anthropic.RateLimitError:
            wait = 60 * (attempt + 1)
            log.warning("Rate limit — waiting %ds", wait)
            time.sleep(wait)
        except anthropic.APIStatusError as e:
            if e.status_code >= 500 and attempt < MAX_RETRIES:
                log.warning("Server error %d — retry %d/%d", e.status_code, attempt + 1, MAX_RETRIES)
                time.sleep(15)
            else:
                raise
        except (httpx.RemoteProtocolError, httpx.ReadError, httpx.ConnectError) as e:
            if attempt < MAX_RETRIES:
                wait = 30 * (attempt + 1)
                log.warning("Connection error (%s) — retry %d/%d in %ds", type(e).__name__, attempt + 1, MAX_RETRIES, wait)
                time.sleep(wait)
            else:
                raise
    raise RuntimeError("Max retries exceeded")


def _extract_json(text: str) -> str:
    """Extract JSON from LLM response, trying fenced blocks first then raw decode."""
    # 1. Explicit ```json fence
    m = re.search(r"```json\s*(.*?)```", text, re.DOTALL)
    if m:
        return m.group(1).strip()
    # 2. Generic ``` fence starting with { or [
    m = re.search(r"```\s*([\[{].*?)```", text, re.DOTALL)
    if m:
        return m.group(1).strip()
    # 3. Raw JSON: find first { and let the decoder determine the correct end
    start = text.find("{")
    if start != -1:
        try:
            obj, _ = json.JSONDecoder().raw_decode(text, start)
            return json.dumps(obj)  # re-serialise to guarantee valid JSON string
        except json.JSONDecodeError:
            pass
    raise ValueError("No JSON found in LLM response")


def _safe_exp_id(exp_id: str) -> str:
    """Make experiment id safe for use as a filename prefix."""
    safe = re.sub(r"\s+", "_", exp_id.strip())
    safe = _UNSAFE_CHARS.sub("_", safe)
    return safe


# ---------------------------------------------------------------------------
# Formulation completeness — detect and repair runs with null quantities

_REFERENCE_PATTERNS = re.compile(
    r"\b(idem|même|identique|cf\.?|voir|same as|see run|like run)\b",
    re.IGNORECASE,
)


def _find_incomplete_runs(doc: dict) -> list[str]:
    """Return run IDs where any formulation component has a null pct_matrix value.

    A component is considered incomplete if:
    - pct_matrix.value is None, AND
    - the component name contains a cross-reference pattern ("idem", "même", etc.)

    We only flag genuine cross-references, not legitimately-absent data.
    """
    incomplete: list[str] = []
    for run in doc.get("runs", []):
        for items in run.get("inputs", {}).values():
            if not isinstance(items, list):
                continue
            for item in items:
                if not isinstance(item, dict):
                    continue
                pm = item.get("pct_matrix", {})
                val = pm.get("value") if isinstance(pm, dict) else None
                name = item.get("component", "")
                if val is None and _REFERENCE_PATTERNS.search(name):
                    incomplete.append(run["id"])
                    break
    return list(dict.fromkeys(incomplete))  # deduplicate, preserve order


def _repair_incomplete_runs(
    doc: dict,
    incomplete_ids: list[str],
    inventory: str,
    client: anthropic.Anthropic,
) -> tuple[dict, _Cost]:
    """Second targeted LLM call (Sonnet 4.6) to fill in missing formulations.

    Uses Sonnet instead of Opus: repair is a narrow, well-defined task.
    Returns (patched_doc, repair_cost).
    """
    repair_cost = _Cost()

    # Build a minimal context: the incomplete runs + their neighbours for reference
    run_ids_set = set(incomplete_ids)
    runs_context = json.dumps(
        [r for r in doc.get("runs", []) if r["id"] in run_ids_set],
        ensure_ascii=False, indent=2
    )

    prompt = f"""\
These runs from a knowledge JSON have incomplete formulations \
(null pct_matrix values where a cross-reference like "idem ESSAI X" was used):

```json
{runs_context}
```

Original file inventory (source of truth):
```
{inventory[:60_000]}
```

Task: for each incomplete run above, find the exact ingredient quantities in the \
inventory and return a JSON array of corrected runs — same structure, same run IDs, \
but with ALL pct_matrix.value fields filled in from the source file.
Return ONLY the JSON array of corrected runs, in ```json ... ``` fences.
Every ingredient must have a numeric pct_matrix.value. \
Never leave null values. Never reference another run.
"""

    try:
        with client.messages.stream(
            model="claude-sonnet-4-6",
            max_tokens=8_000,
            messages=[{"role": "user", "content": prompt}],
        ) as stream:
            text = stream.get_final_text()
            msg = stream.get_final_message()
        repair_cost.add_usage(msg.usage)
        log.info(
            "  repair tokens in=%d out=%d → $%.4f",
            msg.usage.input_tokens, msg.usage.output_tokens, repair_cost.usd,
        )
    except Exception as e:
        log.warning("  repair call failed (%s) — keeping original runs", e)
        return doc, repair_cost

    # Parse the corrected runs
    try:
        json_str = _extract_json(text)
        corrected = json.loads(json_str)
        if not isinstance(corrected, list):
            raise ValueError("Expected a JSON array of runs")
    except Exception as e:
        log.warning("  repair parse failed (%s) — keeping original runs", e)
        return doc, repair_cost

    # Patch the doc: replace each incomplete run with its corrected version
    corrected_by_id = {r["id"]: r for r in corrected}
    patched_runs = []
    fixed = 0
    for run in doc.get("runs", []):
        if run["id"] in corrected_by_id:
            patched_runs.append(corrected_by_id[run["id"]])
            fixed += 1
        else:
            patched_runs.append(run)
    doc["runs"] = patched_runs
    log.info("  repaired %d/%d incomplete run(s)", fixed, len(incomplete_ids))
    return doc, repair_cost


# ---------------------------------------------------------------------------
# Rescue + continuation (for files truncated at 128K output token limit)

def _rescue_from_raw(raw_text: str) -> dict | None:
    """Extract a partial knowledge doc from a truncated LLM response.

    Uses brace/string-aware counting to collect every complete run object
    from the runs array. Returns a doc with complete runs and empty stubs for
    missing sections, or None if nothing recoverable.
    """
    try:
        json_str = _extract_json(raw_text)
        return json.loads(json_str)
    except Exception:
        pass

    json_start = raw_text.find("{")
    if json_start == -1:
        return None
    text = raw_text[json_start:]

    runs_key_pos = text.find('"runs"')
    if runs_key_pos == -1:
        return None
    bracket_pos = text.find("[", runs_key_pos)
    if bracket_pos == -1:
        return None

    # Parse preamble (experiment/targets/glossary/design) by closing the runs array early
    preamble_raw = text[: bracket_pos + 1] + "]\n}"
    try:
        preamble_doc: dict = json.loads(preamble_raw)
    except Exception:
        preamble_doc = {}

    # Walk runs array with brace + string tracking
    complete_runs: list[dict] = []
    pos = bracket_pos + 1
    n = len(text)

    while pos < n:
        while pos < n and text[pos] in " \t\n\r,":
            pos += 1
        if pos >= n or text[pos] != "{":
            break

        depth = 0
        start = pos
        in_string = False
        escape_next = False

        while pos < n:
            ch = text[pos]
            if escape_next:
                escape_next = False
            elif ch == "\\" and in_string:
                escape_next = True
            elif ch == '"':
                in_string = not in_string
            elif not in_string:
                if ch == "{":
                    depth += 1
                elif ch == "}":
                    depth -= 1
                    if depth == 0:
                        try:
                            complete_runs.append(json.loads(text[start : pos + 1]))
                        except Exception:
                            pass
                        pos += 1
                        break
            pos += 1

    if not complete_runs:
        return None

    doc = preamble_doc
    doc["runs"] = complete_runs
    doc.setdefault("failed_runs", [])
    doc.setdefault("derived", [])
    doc.setdefault("observations", {})
    doc.setdefault("not_measured", [])
    doc.setdefault("unused_palette", [])
    return doc


def _build_continuation_message(
    filename: str,
    file_type: str,
    inventory: str,
    existing_run_ids: list[str],
    target_sheet: str | None = None,
) -> str:
    """Prompt asking the LLM for only the parts missing from a truncated extraction."""
    ids_str = ", ".join(f'"{rid}"' for rid in existing_run_ids)
    sheet_note = (
        f"\nTarget sheet (from filename): **{target_sheet}**"
        if target_sheet
        else ""
    )
    return f"""\
## File to extract (CONTINUATION)

Filename: {filename}
File type: {file_type}{sheet_note}

## Context

A previous extraction was truncated at the output token limit.
The following run IDs were already successfully extracted:
[{ids_str}]

## Content dump

```
{inventory}
```

## Task

Generate a PARTIAL knowledge JSON containing ONLY:

1. `"runs"`: every run whose ID is NOT in the already-extracted list above
2. `"failed_runs"`: all failed/aborted runs from the file
3. `"derived"`: all derived/computed quantities (SME, rendement, etc.)
4. `"observations"`: process, sensory, conclusion, next_step
5. `"not_measured"`: analyses not performed

Do NOT repeat runs that are already extracted.
Do NOT include experiment metadata, targets, design, glossary, references, or provenance.

Apply ALL the usual extraction rules (self-contained runs, replicates, units on every measurement, etc.).

Return ONLY the JSON object with those 5 keys, inside ```json ... ``` fences. No explanation.
"""


def _merge_continuation(base_doc: dict, cont_doc: dict) -> dict:
    """Merge continuation doc into base doc.

    Appends new runs; replaces failed_runs, derived, observations, not_measured.
    """
    merged = dict(base_doc)

    existing_ids = {r.get("id") for r in merged.get("runs", [])}
    new_runs = [r for r in cont_doc.get("runs", []) if r.get("id") not in existing_ids]
    merged["runs"] = merged.get("runs", []) + new_runs
    log.info(
        "  continuation merge: +%d new runs → %d total",
        len(new_runs),
        len(merged["runs"]),
    )

    for key in ("failed_runs", "derived", "observations", "not_measured", "unused_palette"):
        if cont_doc.get(key):
            merged[key] = cont_doc[key]

    return merged


# ---------------------------------------------------------------------------
# Main processing loop

def process_file(path: Path, client: anthropic.Anthropic) -> tuple[bool, _Cost]:
    """Returns (success, cost_for_this_file)."""
    stem = path.stem
    out_subdir = OUT_DIR / stem
    out_subdir.mkdir(parents=True, exist_ok=True)
    file_cost = _Cost()

    # Step 1: inventory
    try:
        inventory, target_sheet = get_inventory(path)
    except Exception as e:
        log.error("  INVENTORY FAILED: %s", e)
        return False, file_cost
    log.info(
        "  inventory: %d chars%s",
        len(inventory),
        f" (sheet: {target_sheet!r})" if target_sheet else "",
    )

    # Step 2: LLM extraction
    file_type = (
        "Excel workbook" if path.suffix.lower() in (".xlsx", ".xlsm")
        else "Word document" if path.suffix.lower() == ".docx"
        else "CSV"
    )
    user_msg = _build_user_message(path.name, file_type, inventory, target_sheet)

    try:
        raw, call_cost = _call_llm(user_msg, client)
        file_cost = file_cost + call_cost
    except Exception as e:
        log.error("  LLM FAILED: %s", e)
        return False, file_cost

    # Step 3: parse JSON
    try:
        json_str = _extract_json(raw)
        doc = json.loads(json_str)
    except Exception as e:
        log.error("  JSON PARSE FAILED: %s", e)
        (out_subdir / "_llm_raw_response.txt").write_text(raw, encoding="utf-8")
        log.info("  Raw response saved to %s/_llm_raw_response.txt", stem)
        return False, file_cost

    # Step 3b: detect and repair incomplete formulations
    incomplete = _find_incomplete_runs(doc)
    if incomplete:
        log.warning(
            "  ⚠ %d run(s) with incomplete formulations (cross-references): %s — repairing",
            len(incomplete), incomplete,
        )
        doc, repair_cost = _repair_incomplete_runs(doc, incomplete, inventory, client)
        file_cost = file_cost + repair_cost
    else:
        log.info("  formulation completeness: OK (all runs self-contained)")

    if "experiment" in doc:
        doc["experiment"].setdefault("source_file", path.name)
        if not doc["experiment"].get("sharepoint_url"):
            from src.retrieval.sharepoint_urls import get_url_for_file  # noqa: PLC0415
            url = get_url_for_file(path.name)
            if url:
                doc["experiment"]["sharepoint_url"] = url

    doc.setdefault("provenance", {})
    doc["provenance"]["generated_on"] = TODAY
    doc["provenance"].setdefault("extraction_method", "batch_extract.py LLM")

    exp_id = doc.get("experiment", {}).get("id", stem)
    exp_id_safe = _safe_exp_id(str(exp_id))

    kj_path = out_subdir / f"{exp_id_safe}_knowledge.json"
    kj_path.write_text(json.dumps(doc, ensure_ascii=False, indent=2), encoding="utf-8")
    log.info("  knowledge JSON: %s (%d runs)", kj_path.name, len(doc.get("runs", [])))

    # Step 4: build_kg.py
    result = subprocess.run(
        [sys.executable, str(SCRIPTS_DIR / "build_kg.py"), str(kj_path), "--outdir", str(out_subdir)],
        capture_output=True,
        text=True,
    )
    if result.returncode != 0:
        log.error("  build_kg.py FAILED (rc=%d): %s", result.returncode, result.stderr[:500])
        return False, file_cost

    log.info("  build_kg.py: OK")
    return True, file_cost


def process_file_continuation(path: Path, client: anthropic.Anthropic) -> tuple[bool, _Cost]:
    """Rescue a truncated extraction: recover partial JSON + continuation call for missing runs."""
    stem = path.stem
    out_subdir = OUT_DIR / stem
    raw_path = out_subdir / "_llm_raw_response.txt"
    file_cost = _Cost()

    if not raw_path.exists():
        log.error("  No _llm_raw_response.txt found in %s — cannot rescue", out_subdir)
        return False, file_cost

    raw_text = raw_path.read_text(encoding="utf-8")
    log.info("  rescuing partial JSON from %d-char raw response", len(raw_text))

    base_doc = _rescue_from_raw(raw_text)
    if not base_doc:
        log.error("  could not rescue any data from raw response")
        return False, file_cost

    existing_ids = [r.get("id", "") for r in base_doc.get("runs", [])]
    log.info(
        "  rescued %d complete runs (last: %s)",
        len(existing_ids),
        existing_ids[-1] if existing_ids else "?",
    )

    try:
        inventory, target_sheet = get_inventory(path)
    except Exception as e:
        log.error("  INVENTORY FAILED: %s", e)
        return False, file_cost

    file_type = (
        "Excel workbook" if path.suffix.lower() in (".xlsx", ".xlsm")
        else "Word document" if path.suffix.lower() == ".docx"
        else "CSV"
    )
    user_msg = _build_continuation_message(
        path.name, file_type, inventory, existing_ids, target_sheet
    )

    try:
        raw_cont, call_cost = _call_llm(user_msg, client)
        file_cost = file_cost + call_cost
    except Exception as e:
        log.error("  CONTINUATION LLM FAILED: %s", e)
        return False, file_cost

    try:
        json_str = _extract_json(raw_cont)
        cont_doc = json.loads(json_str)
    except Exception as e:
        log.error("  CONTINUATION JSON PARSE FAILED: %s", e)
        (out_subdir / "_continuation_raw_response.txt").write_text(raw_cont, encoding="utf-8")
        return False, file_cost

    merged_doc = _merge_continuation(base_doc, cont_doc)

    # Detect and repair any incomplete formulations in new runs
    incomplete = _find_incomplete_runs(merged_doc)
    if incomplete:
        log.warning(
            "  ⚠ %d run(s) with incomplete formulations: %s — repairing",
            len(incomplete), incomplete,
        )
        merged_doc, repair_cost = _repair_incomplete_runs(merged_doc, incomplete, inventory, client)
        file_cost = file_cost + repair_cost

    if "experiment" in merged_doc:
        merged_doc["experiment"].setdefault("source_file", path.name)
        if not merged_doc["experiment"].get("sharepoint_url"):
            from src.retrieval.sharepoint_urls import get_url_for_file  # noqa: PLC0415
            url = get_url_for_file(path.name)
            if url:
                merged_doc["experiment"]["sharepoint_url"] = url

    merged_doc.setdefault("provenance", {})
    merged_doc["provenance"]["generated_on"] = TODAY
    merged_doc["provenance"]["extraction_method"] = "batch_extract.py LLM (rescued+continuation)"

    exp_id = merged_doc.get("experiment", {}).get("id", stem)
    exp_id_safe = _safe_exp_id(str(exp_id))

    kj_path = out_subdir / f"{exp_id_safe}_knowledge.json"
    kj_path.write_text(json.dumps(merged_doc, ensure_ascii=False, indent=2), encoding="utf-8")
    log.info(
        "  knowledge JSON: %s (%d runs, %d failed_runs)",
        kj_path.name,
        len(merged_doc.get("runs", [])),
        len(merged_doc.get("failed_runs", [])),
    )

    result = subprocess.run(
        [sys.executable, str(SCRIPTS_DIR / "build_kg.py"), str(kj_path), "--outdir", str(out_subdir)],
        capture_output=True,
        text=True,
    )
    if result.returncode != 0:
        log.error("  build_kg.py FAILED (rc=%d): %s", result.returncode, result.stderr[:500])
        return False, file_cost

    log.info("  build_kg.py: OK")
    return True, file_cost


def main() -> None:
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("--dry-run", action="store_true", help="Show what would be processed")
    ap.add_argument("--force", action="store_true", help="Reprocess even if output exists")
    ap.add_argument("--force-complex", action="store_true",
                    help="Also process level-3 files (normally requires skill). Implies --force.")
    ap.add_argument("--level", type=int, choices=[1, 2, 3], default=0,
                    help="Only process files of this complexity level (0 = 1+2)")
    ap.add_argument("--file", metavar="NAME", help="Process only this filename")
    ap.add_argument("--limit", type=int, default=0, help="Stop after N files (0 = all)")
    ap.add_argument(
        "--rescue",
        action="store_true",
        help=(
            "Rescue a truncated extraction: recover the partial JSON from "
            "_llm_raw_response.txt then call the LLM for the missing runs only. "
            "Use with --file to target a specific file."
        ),
    )
    args = ap.parse_args()

    # --rescue: bypass all skip/level checks, call process_file_continuation
    if args.rescue:
        extensions = {".xlsx", ".xlsm", ".csv", ".docx"}
        candidates_rescue: list[Path] = sorted(
            p for p in BRUT_DIR.iterdir()
            if p.suffix.lower() in extensions
            and not p.name.startswith("~")
            and not p.name.startswith(".")
        )
        if args.file:
            candidates_rescue = [p for p in candidates_rescue if p.name == args.file]
            if not candidates_rescue:
                log.error("File not found: %s", args.file)
                sys.exit(1)
        else:
            # Auto-detect: files that have _llm_raw_response.txt but no knowledge.json
            auto: list[Path] = []
            for p in candidates_rescue:
                out_sub = OUT_DIR / p.stem
                has_raw = (out_sub / "_llm_raw_response.txt").exists()
                has_kg = any(out_sub.glob("*_knowledge.json")) if out_sub.is_dir() else False
                if has_raw and not has_kg:
                    auto.append(p)
            candidates_rescue = auto

        if not candidates_rescue:
            log.info("No truncated files to rescue.")
            sys.exit(0)

        api_key = os.environ.get("ANTHROPIC_API_KEY")
        if not api_key:
            log.error("ANTHROPIC_API_KEY not set")
            sys.exit(1)

        client = anthropic.Anthropic(api_key=api_key)
        limit_r = args.limit if args.limit > 0 else len(candidates_rescue)
        succeeded_r: list[str] = []
        failed_r: list[str] = []
        total_cost_r = _Cost()

        log.info(
            "Pricing: input=$%.2f output=$%.2f cache_write=$%.2f cache_read=$%.2f (per M tokens)",
            _PRICING["input"], _PRICING["output"], _PRICING["cache_write"], _PRICING["cache_read"],
        )

        for i, path in enumerate(candidates_rescue[:limit_r], 1):
            log.info("\n[%d/%d] RESCUE %s", i, min(limit_r, len(candidates_rescue)), path.name)
            ok, file_cost = process_file_continuation(path, client)
            total_cost_r = total_cost_r + file_cost
            log.info(
                "  cost this file: $%.4f  |  running total: $%.4f (%d files)",
                file_cost.usd, total_cost_r.usd, i,
            )
            if ok:
                succeeded_r.append(path.name)
            else:
                failed_r.append(path.name)
            if i < limit_r:
                time.sleep(SLEEP_BETWEEN)

        log.info(
            "\n=== RESCUE COMPLETE — %d OK / %d failed ===",
            len(succeeded_r), len(failed_r),
        )
        if failed_r:
            log.warning("Failed: %s", ", ".join(failed_r))
        log.info(
            "\n💰 COST SUMMARY\n"
            "  Input tokens    : %s\n"
            "  Output tokens   : %s\n"
            "  ─────────────────────────────\n"
            "  TOTAL           : $%.4f USD",
            f"{total_cost_r.input_tokens:>12,}",
            f"{total_cost_r.output_tokens:>12,}",
            total_cost_r.usd,
        )
        log.info(
            "\nNext steps:\n"
            "  1. python -m src.ingest.import_neo4j\n"
            "  2. python -m src.ingest.embed_chunks"
        )
        return

    # --force-complex only unlocks L3 complexity gate — it does NOT bypass the skip set.
    # Use --force explicitly to reprocess already-extracted files.

    skip = _build_skip_set()
    log.info(
        "Skip set: %d source_file matches, %d experiment IDs",
        len(skip.source_files),
        len(skip.exp_ids),
    )

    extensions = {".xlsx", ".xlsm", ".csv", ".docx"}
    candidates: list[Path] = sorted(
        p for p in BRUT_DIR.iterdir()
        if p.suffix.lower() in extensions
        and not p.name.startswith("~")
        and not p.name.startswith(".")
        and ":" not in p.name
    )

    if args.file:
        candidates = [p for p in candidates if p.name == args.file]
        if not candidates:
            log.error("File not found: %s", args.file)
            sys.exit(1)

    to_process: list[Path] = []
    skipped: list[tuple[Path, str]] = []
    deferred_complex: list[Path] = []

    for p in candidates:
        if not args.force:
            if p.name in skip.source_files or p.stem.replace(" ", "_") + p.suffix in skip.source_files:
                skipped.append((p, "source_file match"))
                continue
            if _file_matches_exp_id(p.name, skip.exp_ids):
                skipped.append((p, "experiment ID match"))
                continue
            if _is_already_done(p.stem, args.force):
                skipped.append((p, "output exists"))
                continue

        level = _complexity_level(p.name)
        if args.level and level != args.level:
            skipped.append((p, f"level {level} ≠ requested {args.level}"))
            continue
        if level == 3 and not args.force_complex:
            deferred_complex.append(p)
            continue

        to_process.append(p)

    # Report complexity breakdown
    lvl_counts = {1: 0, 2: 0, 3: 0}
    for p in to_process:
        lvl_counts[_complexity_level(p.name)] += 1

    log.info(
        "To process: %d (L1=%d simple | L2=%d review-needed | L3=%d forced-complex)  "
        "Skipped: %d  Deferred L3: %d",
        len(to_process), lvl_counts[1], lvl_counts[2], lvl_counts[3],
        len(skipped), len(deferred_complex),
    )
    for p, reason in skipped:
        log.debug("  SKIP [%s]: %s", reason, p.name)

    if deferred_complex:
        log.warning(
            "\n⚠  %d file(s) require the interactive skill (instrumental data, replicates):",
            len(deferred_complex),
        )
        for p in deferred_complex:
            log.warning("    → %s", p.name)
        log.warning(
            "  Use the skill in Claude Code for these, or pass --force-complex to batch them\n"
            "  (lower precision — review _validation.md carefully)."
        )

    if args.dry_run:
        log.info("\n--- DRY RUN: would process ---")
        for p in to_process:
            log.info("  [L%d] %s", _complexity_level(p.name), p.name)
        return

    if not to_process:
        log.info("Nothing to process.")
        return

    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        log.error("ANTHROPIC_API_KEY not set")
        sys.exit(1)

    client = anthropic.Anthropic(api_key=api_key)
    limit = args.limit if args.limit > 0 else len(to_process)
    succeeded: list[str] = []
    failed: list[str] = []
    review_needed: list[str] = []
    total_cost = _Cost()

    log.info(
        "Pricing: input=$%.2f output=$%.2f cache_write=$%.2f cache_read=$%.2f (per M tokens)",
        _PRICING["input"], _PRICING["output"], _PRICING["cache_write"], _PRICING["cache_read"],
    )

    for i, path in enumerate(to_process[:limit], 1):
        level = _complexity_level(path.name)
        log.info("\n[%d/%d] [L%d] %s", i, min(limit, len(to_process)), level, path.name)
        out_sub = OUT_DIR / path.stem
        has_raw = (out_sub / "_llm_raw_response.txt").exists()
        has_kg = out_sub.is_dir() and any(out_sub.glob("*_knowledge.json"))
        if has_raw and not has_kg:
            log.info("  detected truncated response — using continuation")
            ok, file_cost = process_file_continuation(path, client)
        else:
            ok, file_cost = process_file(path, client)
        total_cost = total_cost + file_cost
        log.info(
            "  cost this file: $%.4f  |  running total: $%.4f (%d files)",
            file_cost.usd, total_cost.usd, i,
        )
        if ok:
            succeeded.append(path.name)
            if level >= 2:
                review_needed.append(path.name)
        else:
            failed.append(path.name)
        if i < limit:
            time.sleep(SLEEP_BETWEEN)

    log.info("\n=== BATCH COMPLETE — %d OK / %d failed ===", len(succeeded), len(failed))
    if failed:
        log.warning("Failed: %s", ", ".join(failed))

    log.info(
        "\n💰 COST SUMMARY\n"
        "  Input tokens    : %s\n"
        "  Output tokens   : %s\n"
        "  Cache writes    : %s\n"
        "  Cache reads     : %s\n"
        "  ─────────────────────────────\n"
        "  TOTAL           : $%.4f USD",
        f"{total_cost.input_tokens:>12,}",
        f"{total_cost.output_tokens:>12,}",
        f"{total_cost.cache_write_tokens:>12,}",
        f"{total_cost.cache_read_tokens:>12,}",
        total_cost.usd,
    )
    if len(succeeded) > 0:
        log.info("  Average per file: $%.4f", total_cost.usd / len(succeeded))

    if review_needed:
        log.warning(
            "\n📋 %d file(s) need human validation of _validation.md:", len(review_needed)
        )
        for name in review_needed:
            out = OUT_DIR / Path(name).stem
            validation_files = list(out.glob("*_validation.md"))
            val_path = validation_files[0] if validation_files else out / "?_validation.md"
            log.warning("    %s", val_path)

    log.info(
        "\nNext steps:\n"
        "  1. python -m src.ingest.import_neo4j      # import graphs + runs\n"
        "  2. python -m src.ingest.create_indexes    # if new indexes needed\n"
        "  3. python -m src.ingest.embed_chunks      # ← required: Chunk nodes + embeddings"
    )


if __name__ == "__main__":
    main()
