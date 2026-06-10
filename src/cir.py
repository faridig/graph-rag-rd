"""Génération de fiches techniques CIR depuis Neo4j → Claude."""

from __future__ import annotations

import logging
import os
import re
from collections import defaultdict
from collections.abc import Iterator
from dataclasses import dataclass, field

import anthropic
from neo4j import Driver

from src.config import (
    ANTHROPIC_API_KEY,
    CIR_LLM_MODEL,
    NEO4J_PASSWORD,
    NEO4J_URI,
    NEO4J_USER,
)
from src.generation.prompt_cir import (
    CIR_FORMAT,
    SYSTEM_PROMPT_CIR_MUSCLES,
    SYSTEM_PROMPT_CIR_NOUVELLES_VOIES,
    SYSTEM_PROMPT_CIR_PRODUITS,
)

_log = logging.getLogger(__name__)

# Groupements valides tels que définis dans le Répertoire
GROUPEMENTS_VALIDES: list[str] = [
    "Muscles à base de protéines végétales",
    "Produits élaborés à base de muscle végétaux",
    "Nouvelles voies de texturation des protéines végétales",
]

# Le groupement "Nouvelles voies" a ses données sous un autre cir_grouping
# (les runs DST sont classés "Muscles") → requête par chantier
_CHANTIER_DST = "Installation ligne Emincés - Choix de la techno de texturation"

_MAX_CONTEXT_CHARS = 120_000  # ~30k tokens — évite les débordements sur grands groupements

_QUERY_BY_GROUPEMENT = """
MATCH (rep:Run)
WHERE rep.cir_grouping = $cir_grouping
  AND rep.id STARTS WITH "REPERTOIRE"
  AND ($year_prefix IS NULL OR rep.date STARTS WITH $year_prefix)
OPTIONAL MATCH (rep)-[:DETAILS]->(exp:Experiment)
OPTIONAL MATCH (exp)-[:HAS_SUMMARY]->(summary:Chunk)
RETURN rep.id          AS rep_run_id,
       rep.chantier    AS chantier,
       rep.objective   AS objective,
       rep.synthesis   AS synthesis,
       rep.lead        AS lead,
       rep.date        AS date,
       rep.status      AS status,
       exp.id          AS exp_id,
       exp.title       AS exp_title,
       summary.text    AS summary_text
ORDER BY rep.chantier, rep.date
"""

_QUERY_BY_CHANTIER = """
MATCH (rep:Run)
WHERE rep.chantier = $chantier
  AND rep.id STARTS WITH "REPERTOIRE"
  AND ($year_prefix IS NULL OR rep.date STARTS WITH $year_prefix)
OPTIONAL MATCH (rep)-[:DETAILS]->(exp:Experiment)
OPTIONAL MATCH (exp)-[:HAS_SUMMARY]->(summary:Chunk)
RETURN rep.id          AS rep_run_id,
       rep.chantier    AS chantier,
       rep.objective   AS objective,
       rep.synthesis   AS synthesis,
       rep.lead        AS lead,
       rep.date        AS date,
       rep.status      AS status,
       exp.id          AS exp_id,
       exp.title       AS exp_title,
       summary.text    AS summary_text
ORDER BY rep.date
"""

_QUERY_SHAREPOINT_URLS = """
MATCH (e:Experiment)
WHERE e.id IN $exp_ids AND e.sharepoint_url IS NOT NULL
RETURN e.id AS id, e.sharepoint_url AS url
"""


@dataclass
class CirSource:
    run_id: str
    experiment_id: str | None
    sharepoint_url: str | None


@dataclass
class DataQuality:
    runs_total: int
    runs_with_synthesis: int
    runs_with_detailed_data: int
    completeness_pct: int
    warning: str | None


@dataclass
class CirResponse:
    groupement: str
    fiche: str
    data_quality: DataQuality
    sources: list[CirSource] = field(default_factory=list)
    input_tokens: int = 0
    output_tokens: int = 0


@dataclass
class _RunRow:
    rep_run_id: str
    chantier: str | None
    objective: str | None
    synthesis: str | None
    lead: str | None
    date: str | None
    status: str | None
    exp_id: str | None
    exp_title: str | None
    summary_text: str | None


def _fetch_rows(
    driver: Driver, groupement: str, cir_year: int | None = None
) -> list[_RunRow]:
    year_prefix = str(cir_year) if cir_year else None
    with driver.session() as session:
        if groupement == "Nouvelles voies de texturation des protéines végétales":
            records = session.run(
                _QUERY_BY_CHANTIER, chantier=_CHANTIER_DST, year_prefix=year_prefix
            ).data()
        else:
            records = session.run(
                _QUERY_BY_GROUPEMENT, cir_grouping=groupement, year_prefix=year_prefix
            ).data()
    return [
        _RunRow(
            rep_run_id=r["rep_run_id"],
            chantier=r["chantier"],
            objective=r["objective"],
            synthesis=r["synthesis"],
            lead=r["lead"],
            date=r["date"],
            status=r["status"],
            exp_id=r["exp_id"],
            exp_title=r["exp_title"],
            summary_text=r["summary_text"],
        )
        for r in records
    ]


def _fetch_sharepoint_urls(driver: Driver, exp_ids: list[str]) -> dict[str, str]:
    if not exp_ids:
        return {}
    with driver.session() as session:
        records = session.run(_QUERY_SHAREPOINT_URLS, exp_ids=exp_ids).data()
    return {r["id"]: r["url"] for r in records}


def _compute_quality(rows: list[_RunRow]) -> DataQuality:
    total = len(rows)
    with_synthesis = sum(1 for r in rows if r.synthesis)
    with_detail = sum(1 for r in rows if r.summary_text or r.synthesis)
    pct = int(with_detail * 100 / total) if total else 0

    warning: str | None = None
    if total == 0:
        warning = "Aucun essai trouvé pour ce groupement."
    elif with_synthesis < total // 2:
        missing = total - with_synthesis
        warning = (
            f"{missing} runs sans synthèse — "
            "sections Résultats et Nouvelles connaissances partielles."
        )

    return DataQuality(
        runs_total=total,
        runs_with_synthesis=with_synthesis,
        runs_with_detailed_data=with_detail,
        completeness_pct=pct,
        warning=warning,
    )


def _aggregate_by_chantier(rows: list[_RunRow]) -> dict[str, list[_RunRow]]:
    groups: dict[str, list[_RunRow]] = defaultdict(list)
    for r in rows:
        groups[r.chantier or "Sans chantier"].append(r)
    return dict(groups)


def _row_richness(r: _RunRow) -> int:
    """Score de richesse données : summary_text > synthesis > objective only."""
    return (4 if r.summary_text else 0) + (2 if r.synthesis else 0) + (1 if r.objective else 0)


def _format_run(r: _RunRow, urls: dict[str, str]) -> str:
    run_header = f"[{r.rep_run_id}]"
    if r.date:
        run_header += f" {r.date}"
    if r.status:
        run_header += f" [{r.status}]"
    lines = [run_header]
    if r.objective:
        lines.append(f"  Objectif : {r.objective}")
    if r.synthesis:
        lines.append(f"  Synthèse : {r.synthesis}")
    if r.exp_id:
        title_part = f" — {r.exp_title}" if r.exp_title else ""
        lines.append(f"  Essai détaillé : {r.exp_id}{title_part}")
        sp_url = urls.get(r.exp_id)
        if sp_url:
            lines.append(f"  SharePoint : {sp_url}")
    if r.summary_text:
        lines.append(f"  Résultats détaillés :\n{r.summary_text}")
    return "\n".join(lines)


def _format_context(rows: list[_RunRow], urls: dict[str, str]) -> str:
    """Formate le contexte pour Claude.

    Trie les runs par richesse de données (summary > synthesis > objective),
    et tronque si le volume dépasse _MAX_CONTEXT_CHARS pour éviter les
    débordements de contexte sur les grands groupements.
    """
    by_chantier = _aggregate_by_chantier(rows)
    parts: list[str] = []
    total_chars = 0
    skipped = 0

    for chantier, ch_rows in by_chantier.items():
        sorted_rows = sorted(ch_rows, key=_row_richness, reverse=True)
        header = f"=== Chantier : {chantier} ({len(ch_rows)} essais) ==="
        lines = [header]
        for r in sorted_rows:
            block = _format_run(r, urls)
            if total_chars + len(block) > _MAX_CONTEXT_CHARS:
                skipped += 1
                continue
            lines.append(block)
            total_chars += len(block)
        parts.append("\n".join(lines))

    result = "\n\n".join(parts)
    if skipped:
        result += (
            f"\n\n[NOTE CONTEXTE : {skipped} essai(s) omis — volume trop important."
            " Les essais retenus sont ceux avec les données les plus complètes.]"
        )
    return result


def _build_header(groupement: str, rows: list[_RunRow]) -> str:
    dates = [r.date for r in rows if r.date]
    periode = f"{min(dates)} → {max(dates)}" if dates else "N/A"
    leads = sorted({r.lead for r in rows if r.lead})
    return CIR_FORMAT.format(
        groupement=groupement,
        periode=periode,
        leads=", ".join(leads),
        n_essais=len(rows),
    )


def _extract_start_year(rows: list[_RunRow]) -> int | None:
    """Return the year of the earliest dated run, or None if no dates available."""
    dates = [r.date for r in rows if r.date]
    return int(min(dates)[:4]) if dates else None


def get_project_start_year(driver: Driver, groupement: str) -> int | None:
    """Return the year of the earliest run for this groupement (no year filter — full history)."""
    rows = _fetch_rows(driver, groupement, cir_year=None)
    return _extract_start_year(rows)


def _pick_prompt(groupement: str) -> str:
    if groupement == "Muscles à base de protéines végétales":
        return SYSTEM_PROMPT_CIR_MUSCLES
    if groupement == "Nouvelles voies de texturation des protéines végétales":
        return SYSTEM_PROMPT_CIR_NOUVELLES_VOIES
    return SYSTEM_PROMPT_CIR_PRODUITS


def _call_claude(
    client: anthropic.Anthropic,
    groupement: str,
    header: str,
    context: str,
    literature_context: str | None = None,
    start_year: int | None = None,
) -> tuple[str, int, int]:
    prompt = _pick_prompt(groupement)
    if literature_context:
        prompt = f"{prompt}\n\n{literature_context}"
    if start_year:
        prompt = (
            f"{prompt}\n\n"
            f"ANNÉE_DÉMARRAGE : {start_year}\n"
            f"→ Dans §1a, ne citer QUE des publications antérieures à {start_year}. "
            f"Un article publié en {start_year} ou après n'était pas accessible "
            f"au démarrage des travaux et ne peut pas fonder le verrou scientifique."
        )
    user_content = (
        f"En-tête de la fiche (à compléter) :\n{header}\n\n"
        f"Essais R&D disponibles :\n{context}"
    )
    response = client.messages.create(
        model=CIR_LLM_MODEL,
        max_tokens=16000,
        system=prompt,
        messages=[{"role": "user", "content": user_content}],
    )
    return (
        response.content[0].text,
        response.usage.input_tokens,
        response.usage.output_tokens,
    )


def _build_sources(
    rows: list[_RunRow],
    urls: dict[str, str],
) -> list[CirSource]:
    seen: set[str] = set()
    sources: list[CirSource] = []
    for r in rows:
        # Exclude planned runs — not yet realized, not valid evidence in a CIR dossier
        if r.status and "PLANIFIÉ" in r.status.upper():
            continue
        # Exclude RÉPERTOIRE entries with no linked experiment — they have no meaningful
        # label or SharePoint URL and appear as raw composite IDs in the final document.
        if r.exp_id is None:
            continue
        key = r.exp_id
        if key in seen:
            continue
        seen.add(key)
        sources.append(
            CirSource(
                run_id=r.rep_run_id,
                experiment_id=r.exp_id,
                sharepoint_url=urls.get(r.exp_id),
            )
        )
    return sources


def generate_fiche_cir(
    driver: Driver,
    anthropic_client: anthropic.Anthropic,
    groupement: str,
    literature_context: str | None = None,
) -> CirResponse:
    rows = _fetch_rows(driver, groupement)
    quality = _compute_quality(rows)

    exp_ids = [r.exp_id for r in rows if r.exp_id]
    urls = _fetch_sharepoint_urls(driver, exp_ids)
    sources = _build_sources(rows, urls)

    if quality.runs_total < 3:
        fiche = (
            f"⚠ Données insuffisantes pour générer une fiche complète "
            f"({quality.runs_total} essais trouvés, minimum requis : 3).\n"
            f"Les données disponibles sont insuffisantes pour documenter un verrou "
            f"scientifique et une démarche expérimentale au sens Frascati."
        )
        return CirResponse(
            groupement=groupement,
            fiche=fiche,
            data_quality=quality,
            sources=sources,
        )

    header = _build_header(groupement, rows)
    context = _format_context(rows, urls)
    start_year = _extract_start_year(rows)
    fiche, in_tok, out_tok = _call_claude(
        anthropic_client, groupement, header, context, literature_context, start_year
    )

    return CirResponse(
        groupement=groupement,
        fiche=fiche,
        data_quality=quality,
        sources=sources,
        input_tokens=in_tok,
        output_tokens=out_tok,
    )


_MOCK_FICHE = """\
DESCRIPTION DE L'OPÉRATION DE R&D — CIR MESRI

**[MODE TEST — aucun token consommé]**

1. Verrou scientifique et état de l'art
1a. État de l'art
Contenu fictif pour test.

1b. Verrou scientifique
Verrou fictif pour test.

2. Démarche expérimentale
Description fictive des essais.

3. Résultats
3a. Résultats principaux
Résultats fictifs.

3b. Essais non concluants
Aucun essai non concluant dans ce test.

4. Règles opératoires établies
Règles fictives pour test.

5. Perspectives
Perspectives fictives.
"""


def stream_fiche_cir(
    driver: Driver,
    anthropic_client: anthropic.Anthropic,
    groupement: str,
    literature_context: str | None = None,
    cir_year: int | None = None,
) -> Iterator[str | CirResponse]:
    """Yield str tokens during generation, then CirResponse as final item."""
    if os.getenv("CIR_MOCK"):
        yield _MOCK_FICHE
        yield CirResponse(
            groupement=groupement,
            fiche=_MOCK_FICHE,
            data_quality=DataQuality(
                runs_total=0,
                runs_with_synthesis=0,
                runs_with_detailed_data=0,
                completeness_pct=0,
                warning="[MODE TEST]",
            ),
            sources=[],
        )
        return

    rows = _fetch_rows(driver, groupement, cir_year=cir_year)
    quality = _compute_quality(rows)

    exp_ids = [r.exp_id for r in rows if r.exp_id]
    urls = _fetch_sharepoint_urls(driver, exp_ids)
    sources = _build_sources(rows, urls)

    if quality.runs_total < 3:
        fiche = (
            f"⚠ Données insuffisantes pour générer une fiche complète "
            f"({quality.runs_total} essais trouvés, minimum requis : 3).\n"
            "Les données disponibles sont insuffisantes pour documenter un verrou "
            "scientifique et une démarche expérimentale au sens Frascati."
        )
        yield CirResponse(
            groupement=groupement, fiche=fiche, data_quality=quality, sources=sources
        )
        return

    header = _build_header(groupement, rows)
    context = _format_context(rows, urls)
    start_year = _extract_start_year(rows)
    prompt = _pick_prompt(groupement)
    if literature_context:
        prompt = f"{prompt}\n\n{literature_context}"
    if start_year:
        prompt = (
            f"{prompt}\n\n"
            f"ANNÉE_DÉMARRAGE : {start_year}\n"
            f"→ Dans §1a, ne citer QUE des publications antérieures à {start_year}. "
            f"Un article publié en {start_year} ou après n'était pas accessible "
            f"au démarrage des travaux et ne peut pas fonder le verrou scientifique."
        )
    user_content = (
        f"En-tête de la fiche (à compléter) :\n{header}\n\n"
        f"Essais R&D disponibles :\n{context}"
    )

    text_chunks: list[str] = []
    in_tok = out_tok = 0

    with anthropic_client.messages.stream(
        model=CIR_LLM_MODEL,
        max_tokens=16000,
        system=prompt,
        messages=[{"role": "user", "content": user_content}],
    ) as stream:
        for text in stream.text_stream:
            text_chunks.append(text)
            yield text
        final_msg = stream.get_final_message()
        in_tok = final_msg.usage.input_tokens
        out_tok = final_msg.usage.output_tokens

    fiche = "".join(text_chunks)
    yield CirResponse(
        groupement=groupement,
        fiche=fiche,
        data_quality=quality,
        sources=sources,
        input_tokens=in_tok,
        output_tokens=out_tok,
    )


_SEPARATOR_RE = re.compile(r"^━+$")
_MD_BOLD_RE = re.compile(r"\*\*([^*]+)\*\*")


def _clean(text: str) -> str:
    return _MD_BOLD_RE.sub(r"\1", text).strip()


def _add_hyperlink(para: object, url: str, text: str) -> None:
    """Insère un lien hypertexte cliquable dans un paragraphe python-docx."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    r_id = para.part.relate_to(  # type: ignore[attr-defined]
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    r.append(t)
    hyperlink.append(r)
    para._p.append(hyperlink)  # type: ignore[attr-defined]


def export_docx(response: CirResponse, output_path: str) -> None:
    """Write a formatted .docx for the given CirResponse."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()

    for section in doc.sections:
        section.top_margin = section.bottom_margin = Pt(72)
        section.left_margin = section.right_margin = Pt(72)

    lines = response.fiche.split("\n")
    buffer: list[str] = []

    def flush() -> None:
        text = "\n".join(buffer).strip()
        buffer.clear()
        if not text:
            return
        for block in text.split("\n\n"):
            block = _clean(block).strip()
            if not block:
                continue
            if block.startswith("- ") or block.startswith("• "):
                for item in block.splitlines():
                    item = item.lstrip("-• ").strip()
                    if item:
                        p = doc.add_paragraph(style="List Bullet")
                        p.add_run(_clean(item))
            else:
                doc.add_paragraph(_clean(block))

    for line in lines:
        stripped = line.strip()
        clean = _clean(stripped)
        if clean.startswith("#### "):
            flush()
            doc.add_heading(clean[5:], level=3)
        elif clean.startswith("### "):
            flush()
            doc.add_heading(clean[4:], level=2)
        elif clean.startswith("## "):
            flush()
            doc.add_heading(clean[3:], level=1)
        elif clean.startswith("# "):
            flush()
            doc.add_heading(clean[2:], level=1)
        elif _SEPARATOR_RE.match(clean) or clean == "---":
            flush()
        elif clean == "SOURCES":
            flush()
            doc.add_heading("SOURCES", level=2)
        elif any(clean.startswith(k) for k in ("Groupement", "Période", "Leads", "Essais")):
            flush()
            doc.add_paragraph(clean)
        else:
            buffer.append(line)

    flush()

    if response.sources:
        for src in response.sources:
            label = src.experiment_id or src.run_id
            url = src.sharepoint_url or ""
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(label)
            if url:
                p.add_run("  — ")
                _add_hyperlink(p, url, url)

    doc.save(output_path)


def build_cir_clients() -> tuple[Driver, anthropic.Anthropic]:
    from neo4j import GraphDatabase
    driver = GraphDatabase.driver(NEO4J_URI, auth=(NEO4J_USER, NEO4J_PASSWORD))
    client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
    return driver, client
